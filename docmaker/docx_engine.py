#!/usr/bin/env python3
"""
Motor de substituição para o Gerador de Declaração.

Responsável por:
- Carregar o modelo DOCX.
- Aplicar substituições com tentativas em runs e fallback reescrevendo o parágrafo inteiro.
- Remover aspas duplas do resultado.
- Salvar com nome padronizado.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document

TextReplacements = List[Tuple[str, str]]
RegexReplacements = List[Tuple[re.Pattern[str], str]]


@dataclass
class DeclarationData:
    papel: str
    professor: str
    titulo_tcc: str
    autor_tcc: str
    responsavel: str
    cargo: str
    setor: str
    cidade: str
    data_extenso: str


def sanitize_filename(name: str) -> str:
    name = name.strip() or "arquivo"
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    name = name.replace(" ", "_")
    return name[:140]


def _replace_in_paragraph(paragraph, text_repl: TextReplacements, regex_repl: RegexReplacements) -> None:
    # Primeiro: tentar substituir run a run para preservar formatação.
    for run in paragraph.runs:
        run_text = run.text
        for find, repl in text_repl:
            if find in run_text:
                run_text = run_text.replace(find, repl)
        run.text = run_text

    # Fallback: se ainda existe algum alvo ou regex a aplicar, reescreve o parágrafo.
    combined_before = "".join(run.text for run in paragraph.runs)
    combined_after = combined_before
    for find, repl in text_repl:
        combined_after = combined_after.replace(find, repl)
    for pattern, repl in regex_repl:
        combined_after = pattern.sub(repl, combined_after)
    if combined_after != combined_before:
        paragraph.text = combined_after


def _replace_in_container(container, text_repl: TextReplacements, regex_repl: RegexReplacements) -> None:
    for paragraph in container.paragraphs:
        _replace_in_paragraph(paragraph, text_repl, regex_repl)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_container(cell, text_repl, regex_repl)


def _remove_quotes(container) -> None:
    repl = [('"', "")]
    regex_empty: RegexReplacements = []
    _replace_in_container(container, repl, regex_empty)


def _build_replacements(data: DeclarationData) -> Tuple[TextReplacements, RegexReplacements]:
    papel_upper = data.papel.upper()
    papel_lower = data.papel.lower()
    text_repl: TextReplacements = [
        ("DECLARAÇÃO PARECERISTA", f"DECLARAÇÃO {papel_upper}"),
        ("{PAPEL_UPPER}", papel_upper),
        ("{PAPEL}", papel_lower),
        ("parecerista", papel_lower),
        ("{PROFESSOR}", data.professor),
        ("{NOME_PROFESSOR}", data.professor),
        ("{TITULO_TCC}", data.titulo_tcc),
        ("{TÍTULO_TCC}", data.titulo_tcc),
        ("{AUTOR_TCC}", data.autor_tcc),
        ("{RESPONSAVEL}", data.responsavel),
        ("{RESPONSÁVEL}", data.responsavel),
        ("{CARGO_RESPONSAVEL}", data.cargo),
        ("{SETOR}", data.setor),
        ("{CIDADE}", data.cidade),
        ("{DATA}", data.data_extenso),
    ]
    regex_repl: RegexReplacements = [
        (re.compile(r"parecerista", re.IGNORECASE), papel_lower),
    ]
    return text_repl, regex_repl


def apply_declaration(template_path: Path, output_dir: Path, data: DeclarationData) -> Path:
    if not template_path.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {template_path}")

    document = Document(template_path)
    text_repl, regex_repl = _build_replacements(data)

    _replace_in_container(document, text_repl, regex_repl)
    _remove_quotes(document)

    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = sanitize_filename(f"declaracao_{data.papel}_{data.professor}_{timestamp}.docx")
    output_path = output_dir / filename
    document.save(output_path)
    return output_path
