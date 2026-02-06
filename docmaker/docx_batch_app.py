#!/usr/bin/env python3
"""
App Streamlit para gerar DOCX em lote a partir de um modelo e de uma planilha.

Fluxo:
1) Envie o DOCX modelo.
2) Envie a planilha (Excel/CSV) com colunas para cada variável.
3) Preencha o mapeamento: texto exato no DOCX -> coluna da planilha.
4) Gere um DOCX por linha; baixe tudo em ZIP ou use os arquivos salvos em `saida_docx`.
"""
from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import pandas as pd
import streamlit as st
from docx import Document

OUTPUT_DIR = Path("saida_docx")
MAX_SUGGESTIONS = 80
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")
LIST_SEPARATOR = ", "
LIST_CONJUNCTION = "e"


def sanitize_filename(name: str) -> str:
    name = name.strip() or "arquivo"
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    name = name.replace(" ", "_")
    return name[:140]


def build_list_prefix(base_prefix: str, position: int, total: int) -> str:
    prefix = base_prefix or ""
    if total <= 1 or position == 0:
        return prefix
    if position == total - 1:
        needs_space_after = " " if prefix and not prefix.startswith(" ") else ""
        return f" {LIST_CONJUNCTION}{needs_space_after}{prefix}"
    return f"{LIST_SEPARATOR}{prefix}"


def load_spreadsheet(upload) -> pd.DataFrame:
    if upload.name.lower().endswith(".csv"):
        df = pd.read_csv(upload)
    else:
        df = pd.read_excel(upload)
    return df.dropna(how="all")


def extract_text_candidates(doc: Document, limit: int = MAX_SUGGESTIONS) -> List[str]:
    seen = set()
    items: List[str] = []

    def add_text(text: str) -> None:
        val = text.strip()
        if val and val not in seen:
            seen.add(val)
            items.append(val)

    for para in doc.paragraphs:
        add_text(para.text)
        if len(items) >= limit:
            break
    if len(items) < limit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        add_text(para.text)
                        if len(items) >= limit:
                            break
    return items[:limit]


def extract_placeholders_from_docx(buffer: bytes, limit: int = 200) -> List[str]:
    """Varre o XML para capturar {{placeholders}} mesmo que estejam quebrados em runs."""
    with zipfile.ZipFile(io.BytesIO(buffer)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    matches = PLACEHOLDER_RE.findall(xml)
    seen = set()
    result: List[str] = []
    for m in matches:
        key = m.strip()
        if key and key not in seen:
            seen.add(key)
            result.append(f"{{{{{key}}}}}")
        if len(result) >= limit:
            break
    return result


def replace_in_paragraph(paragraph, find: str, replace: str) -> int:
    text = "".join(run.text for run in paragraph.runs)
    if find not in text:
        return 0
    new_text = text.replace(find, replace)
    count = text.count(find)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    return count


def replace_in_container(container, find: str, replace: str) -> int:
    hits = 0
    for paragraph in container.paragraphs:
        hits += replace_in_paragraph(paragraph, find, replace)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                hits += replace_in_container(cell, find, replace)
    return hits


def apply_mapping(doc: Document, mapping: Iterable[Tuple[str, str]]) -> Dict[str, int]:
    stats: Dict[str, int] = {}
    for find, value in mapping:
        stats[find] = replace_in_container(doc, find, value)
    return stats


def remove_quotes(doc: Document) -> None:
    def remove_in_container(container) -> None:
        for paragraph in container.paragraphs:
            text = "".join(run.text for run in paragraph.runs)
            if '"' in text:
                paragraph.text = text.replace('"', "")
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    remove_in_container(cell)

    remove_in_container(doc)


def build_output_name(template: str, row: Dict[str, str], idx: int, primary: str) -> str:
    ctx = {k: str(v) for k, v in row.items()}
    ctx.setdefault("index", idx + 1)
    ctx.setdefault("primary", row.get(primary, "saida"))
    try:
        name = template.format(**ctx)
    except Exception:
        name = f"{idx + 1}_{row.get(primary, 'saida')}"
    return sanitize_filename(name)


def render_app() -> None:
    st.set_page_config(page_title="DOCX em lote", page_icon="📄")
    st.title("Gerador de DOCX com planilha")
    st.write("Mapeie textos do DOCX para colunas do Excel/CSV e gere um arquivo por linha.")

    with st.sidebar:
        st.header("Uploads")
        template_file = st.file_uploader("DOCX modelo", type=["docx"])
        sheet_file = st.file_uploader("Planilha (Excel ou CSV)", type=["xlsx", "xls", "csv"])
        st.caption("Use textos exatos do DOCX como chave de busca.")

    template_bytes = template_file.getvalue() if template_file else None
    df = None
    if sheet_file:
        try:
            df = load_spreadsheet(sheet_file)
        except Exception as exc:  # noqa: BLE001
            st.error(f"Erro ao ler planilha: {exc}")
            return

    suggestions: List[str] = []
    placeholders: List[str] = []
    if template_bytes:
        placeholders = extract_placeholders_from_docx(template_bytes)
        doc = Document(io.BytesIO(template_bytes))
        suggestions = extract_text_candidates(doc)
        if placeholders:
            st.subheader("Placeholders encontrados")
            st.write(", ".join(placeholders))
        elif suggestions:
            st.subheader("Sugestões do DOCX")
            st.write(", ".join(suggestions[:10]) + (" ..." if len(suggestions) > 10 else ""))

    st.subheader("Mapeamento texto -> coluna")
    if "mapping_rows" not in st.session_state:
        st.session_state["mapping_rows"] = [
            {"texto_docx": "", "coluna_planilha": "", "prefixo": "", "grupo": ""} for _ in range(3)
        ]

    column_options = list(df.columns.astype(str)) if df is not None else []

    st.write(
        "Digite o placeholder/texto do DOCX, a coluna da planilha e um texto antes (opcional). "
        "Se quiser, use as sugestões para preencher."
    )
    col1, col2, col3, col4, col5 = st.columns([3, 2, 2, 2, 1])

    # Entrada livre
    placeholder_val = col1.text_input("Placeholder / texto no DOCX", key="ph_text_input")
    column_val = col2.text_input("Coluna do Excel", key="col_text_input")
    prefix_val = col3.text_input("Texto antes (prefixo)", key="prefix_text_input")
    group_val = col4.text_input("Grupo (lista)", key="group_text_input")

    # Sugestões opcionais
    sel_text = col1.selectbox(
        "Sugestões do DOCX",
        options=[""] + (placeholders or suggestions),
        key="sel_textbox",
        help="Use placeholders {{campo}} se existirem; senão use textos detectados."
    )
    if sel_text:
        placeholder_val = sel_text
        st.session_state["ph_text_input"] = sel_text

    sel_col = col2.selectbox("Colunas detectadas", options=[""] + column_options, key="sel_colbox")
    if sel_col:
        column_val = sel_col
        st.session_state["col_text_input"] = sel_col

    if col5.button("Adicionar", use_container_width=True):
        if placeholder_val and column_val:
            st.session_state["mapping_rows"].append(
                {
                    "texto_docx": placeholder_val,
                    "coluna_planilha": column_val,
                    "prefixo": prefix_val,
                    "grupo": group_val,
                }
            )
        else:
            st.warning("Preencha placeholder/texto e coluna.")

    edited = st.data_editor(
        st.session_state["mapping_rows"],
        num_rows="dynamic",
        key="mapping_editor",
        column_config={
            "texto_docx": st.column_config.TextColumn("Texto no DOCX", help="Trecho exato a ser substituído."),
            "coluna_planilha": st.column_config.SelectboxColumn(
                "Coluna do Excel", options=column_options, help="Coluna cujo valor entra no lugar do texto."
            ),
            "prefixo": st.column_config.TextColumn(
                "Texto antes",
                help="Texto fixo antes do valor. Se o valor estiver vazio, nada será exibido.",
            ),
            "grupo": st.column_config.TextColumn(
                "Grupo (lista)",
                help="Use o mesmo grupo para aplicar ', ' e ' e ' automaticamente.",
            ),
        },
        use_container_width=True,
    )
    st.session_state["mapping_rows"] = edited

    name_template = st.text_input("Template do nome do arquivo", "{index:03d}_{primary}.docx")
    primary_col = None
    if column_options:
        primary_col = st.selectbox("Coluna usada no nome do arquivo (primary)", options=column_options, index=0)

    if st.button("Gerar documentos", type="primary"):
        if not template_bytes:
            st.error("Envie o DOCX modelo.")
            return
        if df is None or df.empty:
            st.error("Envie a planilha com dados.")
            return

        mapping = []
        for row in edited:
            find = (row.get("texto_docx") or "").strip()
            col = (row.get("coluna_planilha") or "").strip()
            prefix = row.get("prefixo") or ""
            group = (row.get("grupo") or "").strip()
            if find and col:
                mapping.append((find, col, prefix, group))
        if not mapping:
            st.error("Preencha ao menos um mapeamento.")
            return

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        results: List[Tuple[str, Dict[str, int], bytes]] = []
        for idx, (_, data_row) in enumerate(df.iterrows()):
            doc = Document(io.BytesIO(template_bytes))
            items = []
            for map_idx, (find, col, prefix, group) in enumerate(mapping):
                value = data_row.get(col, "")
                if pd.isna(value):
                    value = ""
                value_str = str(value)
                items.append(
                    {
                        "index": map_idx,
                        "find": find,
                        "prefix": prefix or "",
                        "group": group,
                        "value": value_str,
                        "has_value": bool(value_str.strip()),
                    }
                )

            seen_by_group: Dict[str, set] = {}
            for item in items:
                if not item["group"] or not item["has_value"]:
                    continue
                normalized = item["value"].strip().lower()
                seen = seen_by_group.setdefault(item["group"], set())
                if normalized in seen:
                    item["has_value"] = False
                else:
                    seen.add(normalized)

            group_positions: Dict[int, Tuple[int, int]] = {}
            groups: Dict[str, List[dict]] = {}
            for item in items:
                if not item["group"]:
                    continue
                groups.setdefault(item["group"], []).append(item)
            for group_items in groups.values():
                present = [item for item in group_items if item["has_value"]]
                present.sort(key=lambda x: x["index"])
                total = len(present)
                for pos, item in enumerate(present):
                    group_positions[item["index"]] = (pos, total)

            replacements: List[Tuple[str, str]] = []
            for item in items:
                if not item["has_value"]:
                    replacements.append((item["find"], ""))
                    continue
                prefix = item["prefix"]
                if item["group"] and item["index"] in group_positions:
                    pos, total = group_positions[item["index"]]
                    prefix = build_list_prefix(prefix, pos, total)
                replacements.append((item["find"], f"{prefix}{item['value']}"))
            stats = apply_mapping(doc, replacements)
            remove_quotes(doc)
            filename = build_output_name(name_template, dict(data_row), idx, primary_col or mapping[0][1])

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            (OUTPUT_DIR / filename).write_bytes(buffer.getvalue())
            results.append((filename, stats, buffer.getvalue()))

        st.success(f"Gerados {len(results)} arquivo(s). Salvos em `{OUTPUT_DIR}`.")
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for filename, _, content in results:
                zf.writestr(filename, content)
        zip_buffer.seek(0)
        st.download_button("Baixar todos em ZIP", zip_buffer, file_name="gerados.zip")

        with st.expander("Ocorrências por arquivo"):
            for filename, stats, _ in results:
                st.write(f"`{filename}` -> {stats}")


if __name__ == "__main__":
    render_app()
